from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
ASSET_DIR = ROOT / "build" / "report_assets"
OUT = REPORT_DIR / "VeeR-EH2双DDR4系统设计与验证报告_GCCLKT0最终版.docx"

PRESET = "standard_business_brief"
COVER_PATTERN = "editorial_cover"

BLUE = "0F4C81"
DARK_BLUE = "17324D"
MID_BLUE = "2B6F9F"
TEAL = "007C91"
GREEN = "2E7D32"
AMBER = "B26A00"
RED = "B3261E"
LIGHT_BLUE = "EAF3F8"
LIGHT_GREEN = "EAF4EA"
LIGHT_AMBER = "FFF3DD"
LIGHT_RED = "FCE8E6"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D7DDE3"
TEXT = "253746"
MUTED = "5E6C78"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths: Sequence[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header(row) -> None:
    repeat_table_header(row)


def set_run_font(run, size: float | None = None, bold: bool | None = None,
                 color: str | None = None, latin: str = "Calibri",
                 east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 28, WHITE, 0, 10),
        ("Subtitle", 13, "D9EAF5", 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.08

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Code Box" not in styles:
        code_style = styles.add_style("Code Box", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Box"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "Small Note" not in styles:
        note = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        note = styles["Small Note"]
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(8.5)
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.space_after = Pt(4)
    note.paragraph_format.line_spacing = 1.05


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("VeeR-EH2  |  双 DDR4 硬件系统")
    set_run_font(run, 8.5, True, BLUE)
    p_pr = p._p.get_or_add_pPr()
    bottom_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), MID_GREY)
    bottom_border.append(bottom)
    p_pr.append(bottom_border)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("工程技术报告  •  2026-07-24  •  ")
    set_run_font(fr, 8, False, MUTED)
    add_page_number(fp)


def add_label_value(paragraph, label: str, value: str) -> None:
    r = paragraph.add_run(label)
    set_run_font(r, 10, True, DARK_BLUE)
    r = paragraph.add_run(value)
    set_run_font(r, 10, False, TEXT)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.page_break_before = False


def add_paragraph(doc: Document, text: str = "", style: str | None = None,
                  bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        add_paragraph(doc, item, style="List Bullet")


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    numbering = doc.part.numbering_part.element
    style_num_pr = doc.styles["List Number"]._element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    base_num = next(
        node for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_num_id = max(existing_ids, default=0) + 1
    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_num_id))
    new_num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)

    for item in items:
        p = doc.add_paragraph(style="List Number")
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(new_num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        p._p.get_or_add_pPr().append(num_pr)
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], header_fill: str = LIGHT_GREY,
              first_col_bold: bool = False) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        set_cell_shading(hdr.cells[i], header_fill)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, 9, True, DARK_BLUE)
    set_repeat_header(hdr)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            p = row.cells[i].paragraphs[0]
            r = p.add_run(str(text))
            set_run_font(r, 8.7, bool(first_col_bold and i == 0), TEXT)
    set_table_fixed(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, title: str, body: str, fill: str, accent: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [9200])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, 10.2, True, accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, 9.5, False, TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_box(doc: Document, lines: Sequence[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_fixed(table, [9200])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FA")
    p = cell.paragraphs[0]
    p.style = doc.styles["Code Box"]
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, 8.5, False, DARK_BLUE, latin="Consolas")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def load_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size, index=0)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, title, body_lines, title_font, body_font, radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = xy
    draw.text((x1 + 18, y1 + 15), title, font=title_font, fill="#17324D")
    y = y1 + 58
    for line in body_lines:
        draw.text((x1 + 18, y), line, font=body_font, fill="#253746")
        y += 31


def arrow(draw, start, end, color="#2B6F9F", width=6):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 18 * sign, y2 - 11), (x2 - 18 * sign, y2 + 11)]
    else:
        sign = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 11, y2 - 18 * sign), (x2 + 11, y2 - 18 * sign)]
    draw.polygon(pts, fill=color)


def make_architecture_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(image)
    title_font = load_font(30, True)
    body_font = load_font(24)
    small_font = load_font(21)
    draw.text((70, 35), "双 DDR4 系统结构与时钟域", font=load_font(38, True), fill="#0F4C81")

    rounded_box(draw, (70, 150, 390, 400), "#EAF3F8", "#2B6F9F", "VeeR-EH2", ["IFU AXI 64 bit", "LSU AXI 64 bit", "DMA AXI 初始化口", "core_clk = 50 MHz"], title_font, body_font)
    rounded_box(draw, (520, 120, 880, 330), "#F2F4F7", "#607D8B", "DDR4-1 前端", ["ATG0 32→512", "IFU 64→512", "AXI CDC + 所有权选择"], title_font, body_font)
    rounded_box(draw, (520, 410, 880, 660), "#F2F4F7", "#607D8B", "DDR4-2 前端", ["ATG1 32→512", "LSU 64→512", "终态结果读取器", "AXI CDC + 所有权选择"], title_font, body_font)
    rounded_box(draw, (1020, 120, 1350, 330), "#EAF4EA", "#2E7D32", "MIG DDR4-1", ["AXI 512 bit", "UI = 266.525 MHz", "ECC / x72"], title_font, body_font)
    rounded_box(draw, (1020, 410, 1350, 620), "#EAF4EA", "#2E7D32", "MIG DDR4-2", ["AXI 512 bit", "UI = 266.525 MHz", "ECC / x72"], title_font, body_font)
    rounded_box(draw, (1480, 120, 1730, 330), "#FFF3DD", "#B26A00", "DDR4-1", ["程序存储", "8 GB 级"], title_font, body_font)
    rounded_box(draw, (1480, 410, 1730, 620), "#FFF3DD", "#B26A00", "DDR4-2", ["数据存储", "8 GB 级"], title_font, body_font)

    arrow(draw, (390, 235), (520, 220))
    arrow(draw, (390, 320), (520, 500))
    arrow(draw, (880, 225), (1020, 225))
    arrow(draw, (880, 515), (1020, 515))
    arrow(draw, (1350, 225), (1480, 225))
    arrow(draw, (1350, 515), (1480, 515))

    draw.text((420, 175), "IFU", font=small_font, fill="#0F4C81")
    draw.text((420, 410), "LSU", font=small_font, fill="#0F4C81")
    draw.text((95, 735), "外部差分时钟", font=title_font, fill="#17324D")
    clock_rows = [
        ("GCLK2 Top", "50.000 MHz", "EH2"),
        ("GCLK1 Top", "100.000 MHz", "ATG"),
        ("GCLK3 Top", "76.150 MHz", "MIG DDR4-1"),
        ("GCLK3 Bottom", "76.150 MHz", "MIG DDR4-2"),
    ]
    x = 95
    for name, freq, target in clock_rows:
        rounded_box(draw, (x, 800, x + 370, 1010), "#F7F9FA", "#D7DDE3", name, [freq, target], body_font, small_font, radius=14)
        x += 420
    image.save(path, dpi=(180, 180))


def make_boot_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 790), "white")
    draw = ImageDraw.Draw(image)
    title = load_font(37, True)
    step_font = load_font(25, True)
    body_font = load_font(20)
    draw.text((70, 35), "上电、初始化、总线交接与结果验证时序", font=title, fill="#0F4C81")
    steps = [
        ("1", "系统复位", ["EH2 停止", "ATG/MIG 复位"]),
        ("2", "MIG 校准", ["两路 init_calib_complete", "同步到控制域"]),
        ("3", "双 ATG 装载", ["DDR1 写程序", "DDR2 写初始数据"]),
        ("4", "永久交接", ["done_latched 置位", "ATG 永久保持复位"]),
        ("5", "TCM 清零", ["EH2 处于 debug halt", "DMA AXI 写零并生成 ECC"]),
        ("6", "运行与检查", ["debug run → PC=0", "读回 0x1000C = 0x1BC"]),
    ]
    colors = ["#F2F4F7", "#EAF3F8", "#FFF3DD", "#EAF4EA", "#EAF3F8", "#EAF4EA"]
    outlines = ["#607D8B", "#2B6F9F", "#B26A00", "#2E7D32", "#2B6F9F", "#2E7D32"]
    x = 50
    for idx, (num, name, body) in enumerate(steps):
        box = (x, 185, x + 255, 570)
        draw.rounded_rectangle(box, radius=20, fill=colors[idx], outline=outlines[idx], width=3)
        draw.ellipse((x + 91, 205, x + 165, 279), fill=outlines[idx])
        nbox = draw.textbbox((0, 0), num, font=step_font)
        nw = nbox[2] - nbox[0]
        nh = nbox[3] - nbox[1]
        draw.text((x + 128 - nw / 2, 242 - nh / 2 - 3), num, font=step_font, fill="white")
        draw.text((x + 28, 305), name, font=step_font, fill="#17324D")
        y = 365
        for line in body:
            draw.text((x + 20, y), line, font=body_font, fill="#253746")
            y += 48
        if idx < len(steps) - 1:
            arrow(draw, (x + 255, 380), (x + 295, 380), width=5)
        x += 295

    draw.rounded_rectangle((270, 650, 1530, 738), radius=15, fill="#FCE8E6", outline="#B3261E", width=2)
    draw.text((300, 672), "关键保证：ATG done 锁存后，其 resetn 一直为 0；主数据通路的 mux 只会从 ATG 切到 EH2，不会返回。", font=body_font, fill="#8A1C16")
    image.save(path, dpi=(180, 180))


def add_cover(doc: Document) -> None:
    banner = doc.add_table(rows=1, cols=1)
    banner.style = "Table Grid"
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(banner, [9360])
    cell = banner.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_margins(cell, top=520, bottom=520, start=420, end=420)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.style = doc.styles["Title"]
    r = p.add_run("VeeR-EH2 双 DDR4 系统")
    set_run_font(r, 28, True, WHITE)
    p = cell.add_paragraph()
    p.style = doc.styles["Subtitle"]
    r = p.add_run("设计、初始化、总线交接与验证报告")
    set_run_font(r, 15, False, "D9EAF5")
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("目标器件  xcvu19p_CIV-fsva3824-1-e")
    set_run_font(r, 11.5, True, WHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(38)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("工程结论")
    set_run_font(r, 13, True, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("双 MIG、双 ATG、EH2 硬件初始化、AXI 跨时钟/宽度转换、永久总线交接、结果检查及 LED 指示均已集成。前仿、综合、布局布线、时序分析与比特流生成完成。")
    set_run_font(r, 14, False, DARK_BLUE)

    add_table(doc, ["项目", "状态"], [
        ("系统前仿", "通过：DDR4-2 最终值 0x000001BC，LED=1111"),
        ("完整 TCM 清零仿真", "通过：16,384 次有序 64 位写零"),
        ("布局布线后时序", "通过：WNS +0.202 ns，TNS 0"),
        ("比特流", "已生成：eh2_dual_ddr_top.bit"),
        ("物理板卡实测", "尚未执行；需下载比特流后完成上板确认"),
    ], [2400, 6960], header_fill="D9EAF5", first_col_bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    rows = [
        ("报告版本", "V1.0", "生成日期", "2026-07-24"),
        ("设计工具", "Vivado 2023.2", "工程目录", str(ROOT)),
    ]
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            if col_idx in (0, 2):
                set_cell_shading(info.cell(row_idx, col_idx), LIGHT_GREY)
            p = info.cell(row_idx, col_idx).paragraphs[0]
            r = p.add_run(value)
            set_run_font(r, 8.3, col_idx in (0, 2), DARK_BLUE if col_idx in (0, 2) else TEXT)
    set_table_fixed(info, [1450, 1800, 1450, 4660])


def build_report() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    architecture_png = ASSET_DIR / "dual_ddr_architecture.png"
    boot_png = ASSET_DIR / "boot_and_handoff.png"
    make_architecture_diagram(architecture_png)
    make_boot_diagram(boot_png)

    doc = Document()
    style_doc(doc)
    configure_sections(doc)
    props = doc.core_properties
    props.title = "VeeR-EH2 双 DDR4 系统设计与验证报告"
    props.subject = "双 MIG、ATG 初始化、EH2 硬件初始化与 FPGA 实现验证"
    props.author = "Codex / FPGA 工程工作区"
    props.keywords = "VeeR-EH2, DDR4, MIG, AXI, ATG, ECC, FPGA"
    props.comments = f"Preset={PRESET}; cover={COVER_PATTERN}"

    add_cover(doc)
    add_page_break(doc)

    add_heading(doc, "1. 执行摘要", 1)
    add_callout(
        doc,
        "最终状态：实现完成，静态检查通过",
        "工程已完成系统级前仿、完整 DCCM/ICCM 清零仿真、综合后功能仿真、综合、布局布线、时序分析和比特流生成。布局布线后 WNS 为 +0.202 ns、TNS 为 0；44 项总线偏斜检查全部满足；路由错误为 0。",
        LIGHT_GREEN,
        GREEN,
    )
    add_paragraph(doc, "本设计把 VeeR-EH2 的取指 AXI 端口连接到 DDR4-1，把 LSU 访存 AXI 端口连接到 DDR4-2。两块 DDR4 均采用独立 MIG，MIG 参数与 MAC 参考工程保持一致，并按 VeriTiger-V19P-A14 用户手册完成第二组 DDR4 管脚与差分时钟约束。")
    add_paragraph(doc, "DDR4-1 由一次性 ATG 写入可执行程序，DDR4-2 由一次性 ATG 写入程序所需初始数据。每个 ATG 完成后，其完成状态被锁存，ATG 复位被永久拉低，同时 AXI 所有权切换到相应 EH2 主机；只有板级系统复位才能重新进入 ATG 阶段。")
    add_paragraph(doc, "EH2 在进入正常取指前保持 debug halt。硬件初始化模块通过 EH2 的 DMA AXI 从端口，对 DCCM 64 KiB 与 ICCM 64 KiB 各执行 8192 次 64 位写零；写入经过 EH2 正常存储路径，由内部存储体生成正确 39 位 SECDED ECC。全部响应正确后才发出 debug run 请求，使处理器从复位向量 0 开始执行。")

    add_heading(doc, "1.1 交付范围", 2)
    add_bullets(doc, [
        "可综合 RTL：双 MIG 连接、ATG 前端、AXI 所有权选择、跨时钟与宽度转换、EH2 包装、硬件 TCM 初始化、结果读取与 LED 状态逻辑。",
        "约束：双 DDR4 管脚、四组差分输入时钟、复位开关、LED、时钟分组与 CDC 相关属性。",
        "存储初始化：DDR 兼容程序映像与两组 ATG 命令/地址/数据/掩码 COE。",
        "验证与实现：系统前仿、完整 TCM 清零仿真、综合和实现报告、实现后 DCP 及最终 BIT。",
        "本报告：模块职责、时钟频率、初始化机制、总线所有权、程序兼容性改动和实现结果。",
    ])
    add_callout(doc, "验证边界", "当前结论基于 RTL 仿真和 Vivado 静态实现结果。尚未把比特流下载到实体板卡，因此 LED 物理点亮、DDR 温度裕量和板级信号完整性仍需上板验证。", LIGHT_AMBER, AMBER)

    add_page_break(doc)
    add_heading(doc, "2. 系统架构", 1)
    add_paragraph(doc, "系统使用四个相互独立的外部时钟入口。两块 DDR4 分别使用板卡上下区域各自的 GCLK3 差分对；它们频率相同但不是同一物理时钟网络。EH2 使用 GCLK2 Top 的 50 MHz 差分时钟，ATG 使用 GCLK1 Top 的 100 MHz 差分时钟。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(architecture_png), width=Inches(6.45))
    add_caption(doc, "图 2-1  双 DDR4 系统结构、数据路径与时钟域")

    add_heading(doc, "2.1 端口到 DDR 的映射", 2)
    add_table(doc, ["EH2 / 主机端口", "目标", "用途", "主数据宽度"], [
        ("IFU AXI", "MIG DDR4-1", "从地址 0 取指", "64 bit → 512 bit"),
        ("LSU AXI", "MIG DDR4-2", "普通 load/store 数据访问", "64 bit → 512 bit"),
        ("EH2 DMA AXI 从端口", "DCCM / ICCM", "复位释放前的架构启动阶段写零和 ECC 初始化", "64 bit"),
        ("ATG0", "MIG DDR4-1", "一次性写入程序映像", "32 bit → 512 bit"),
        ("ATG1", "MIG DDR4-2", "一次性写入数据初值", "32 bit → 512 bit"),
        ("结果读取器", "MIG DDR4-2", "终态读取 0x1000C 并判断 0x1BC", "512 bit beat"),
    ], [2100, 1900, 3600, 1760], first_col_bold=True)

    add_heading(doc, "2.2 模块职责与工作时钟", 2)
    add_table(doc, ["模块 / 实例", "系统责任", "主要时钟"], [
        ("eh2_dual_ddr_top", "顶层连接、复位序列、状态汇总和三颗 LED 输出", "50 / 100 / 266.525 MHz 多域"),
        ("eh2_core_wrapper_hw", "封装 EH2、暴露 IFU/LSU/DMA/debug 接口并设置 rst_vec=0", "core_clk 50 MHz"),
        ("eh2_hw_init", "在 debug halt 状态经 DMA AXI 清零 64 KiB DCCM 与 64 KiB ICCM", "core_clk 50 MHz"),
        ("axi64_to_512_cdc ×2", "IFU/LSU 的 Xilinx AXI Clock Converter + Data Width Converter", "50 MHz ↔ MIG UI 266.525 MHz"),
        ("axi32_to_512_cdc ×2", "ATG 的 Xilinx AXI Clock Converter + Data Width Converter", "100 MHz ↔ MIG UI 266.525 MHz"),
        ("axi_owner_mux2 ×2", "在同步后的 sticky done 控制下把主通路从 ATG 切到 EH2", "各 MIG UI 266.525 MHz"),
        ("ddr_result_checker", "检测终态写响应后读取 DDR4-2，并锁存 pass/fail", "DDR4-2 MIG UI 266.525 MHz"),
        ("MIG DDR4 ×2", "DDR4 PHY、校准、ECC 与 512 位 AXI 用户接口", "GCLK3 76.15 MHz；UI 266.525 MHz"),
        ("ATG ×2", "上电后一次性生成程序/数据 AXI 写事务", "atg_clk 100 MHz"),
    ], [2350, 4720, 2290], first_col_bold=True)

    add_page_break(doc)
    add_heading(doc, "3. DDR4 MIG 配置与板级约束", 1)
    add_heading(doc, "3.1 两路 MIG 的共同参数", 2)
    add_table(doc, ["配置项", "取值", "设计含义"], [
        ("存储器类型", "DDR4 SODIMM，single slot", "对应板载两组 SODIMM 接口"),
        ("器件模型", "MTA9ASF1G72HZ-2G6", "x72，含 8 位 ECC 数据线"),
        ("电压", "1.2 V", "DDR4 标准工作电压"),
        ("数据宽度 / ECC", "72 bit / ECC enabled", "64 位数据 + 8 位 ECC"),
        ("AXI 用户接口", "ADDR 33、DATA 512、ID 4", "允许 narrow burst"),
        ("输入时钟周期", "13132 ps", "外部 GCLK3 = 76.15 MHz"),
        ("DDR 周期", "938 ps", "约 1066 MHz CK / 2133 MT/s"),
        ("PHY / UI", "4:1 / 266.525 MHz", "MIG AXI 用户时钟"),
        ("读写时序", "CL 15 / CWL 11", "与参考工程配置一致"),
    ], [2450, 2850, 4060], first_col_bold=True)

    add_paragraph(doc, "DDR4-2 的 MIG 由 DDR4-1 的配置复制并修改为第二组物理接口；电气参数、AXI 参数、ECC 和时序参数保持一致。两路 MIG 各自独立执行 PHY 校准，各自输出 init_calib_complete 和 UI 时钟。")

    add_heading(doc, "3.2 外部时钟、复位与 LED 管脚", 2)
    add_table(doc, ["信号", "管脚", "频率 / 电平", "用途"], [
        ("c0_sys_clk_p/n", "BN26 / BP26", "76.15 MHz LVDS", "DDR4-1，GCLK3 Top"),
        ("c1_sys_clk_p/n", "F32 / E32", "76.15 MHz LVDS", "DDR4-2，GCLK3 Bottom"),
        ("core_clk_p/n", "BY44 / CA44", "50.00 MHz LVDS", "EH2，GCLK0 Top (GCCLKT0)"),
        ("atg_clk_p/n", "BN55 / BP55", "100.00 MHz LVDS", "双 ATG，GCLK1 Top"),
        ("system_resetn", "BU21", "LVCMOS12，高有效释放", "全系统板级复位开关"),
        ("aux_resetn", "BU28", "LVCMOS12，高有效释放", "辅助复位开关"),
        ("led[0]", "BE22", "LVCMOS12，高亮", "两路 ATG 均完成且无错误"),
        ("led[1]", "BG23", "LVCMOS12，高亮", "检测到 EH2 IFU 取指地址握手"),
        ("led[2]", "BJ20", "LVCMOS12，高亮", "检测到 EH2 LSU 写操作"),
        ("led[3]", "BN19", "LVCMOS12，高亮", "TCM 初始化成功且 DDR 终值检查通过"),
    ], [2050, 1800, 2300, 3210], first_col_bold=True)
    add_callout(doc, "GCLK3 使用说明", "两块 DDR 都使用名为 GCLK3 的板级资源，但 DDR4-1 使用 Top 的 BN26/BP26，DDR4-2 使用 Bottom 的 F32/E32；它们是两对独立差分输入，不共用一根时钟线。", LIGHT_BLUE, BLUE)

    add_heading(doc, "3.3 DDR 管脚与时序约束策略", 2)
    add_paragraph(doc, "完整管脚约束位于 constraints/eh2_dual_ddr_v19p.xdc。每路接口分别约束 72 根 DQ、9 对 DQS、地址/Bank/Bank Group、差分 CK、CKE、CS、ODT、ACT、PAR、ALERT 和 RESET。第二路管脚映射依据 VeriTiger-V19P-A14 用户手册的 DDR4-2 连接表建立；第一路保持与 MAC 参考工程一致。")
    add_bullets(doc, [
        "core_clk、atg_clk、c0_sys_clk 和 c1_sys_clk 均创建为独立主时钟。",
        "两个 MIG 生成的 UI 时钟由 IP 自带约束派生，布局布线后频率均为 266.525 MHz。",
        "core、ATG、MIG0 和 MIG1 时钟域之间设置异步时钟组；多位 AXI 数据通过 Xilinx AXI Clock Converter 处理。",
        "ATG done、校准完成等单比特状态采用两级 ASYNC_REG 同步；六个状态分别同步后再组合，避免组合信号跨域。",
    ])

    add_page_break(doc)
    add_heading(doc, "4. AXI 互连、跨时钟与宽度转换", 1)
    add_paragraph(doc, "互连结构按 MAC 工程的做法使用 Xilinx 标准 AXI IP，而不是自写异步 FIFO。ATG 和 EH2 的不同数据宽度先由 AXI Data Width Converter 转换为 MIG 所需的 512 bit，再由 AXI Clock Converter 完成时钟域转换。")
    add_table(doc, ["通路", "源端", "转换", "目标端"], [
        ("ATG0 → DDR4-1", "32 bit @ 100 MHz", "32→512 + async CDC（3 stages）", "512 bit @ 266.525 MHz"),
        ("ATG1 → DDR4-2", "32 bit @ 100 MHz", "32→512 + async CDC（3 stages）", "512 bit @ 266.525 MHz"),
        ("IFU → DDR4-1", "64 bit @ 50 MHz", "64→512 + async CDC（3 stages）", "512 bit @ 266.525 MHz"),
        ("LSU → DDR4-2", "64 bit @ 50 MHz", "64→512 + async CDC（3 stages）", "512 bit @ 266.525 MHz"),
    ], [2050, 2100, 3030, 2180], first_col_bold=True)
    add_paragraph(doc, "MIG AXI 参数统一为 33 位地址、4 位 ID、512 位数据，并允许 narrow burst。EH2 IFU/LSU 原生 ID 和响应在包装层保留，转换 IP 负责宽度拆装、WSTRB 映射、burst 处理和返回通道背压。")

    add_heading(doc, "4.1 CDC 处理原则", 2)
    add_bullets(doc, [
        "AXI 五个通道的多位数据与握手由 AXI Clock Converter 异步 FIFO 处理。",
        "ATG done_latched、MIG 校准完成、错误状态等控制位均先独立两级同步，再在目标域组合。",
        "同步寄存器标记 ASYNC_REG，使实现工具采用适合亚稳态收敛的布局和 CDC 识别。",
        "复位在各时钟域中异步置位、同步释放；实现报告识别到的复位同步器均为信息级结构。",
    ])

    add_heading(doc, "4.2 地址与写选通示例", 2)
    add_paragraph(doc, "最终 32 位 store 的目标为 0x1000C。在 EH2 64 位 LSU 接口上，它使用对齐地址 0x10008，并通过 WSTRB[7:4] 写入 WDATA[63:32]。经过 64→512 转换后，MIG 看到 0x10000 对齐的 512 位 beat；结果读取器从返回数据的 [127:96] 提取最终字。")
    add_code_box(doc, [
        "EH2 store target : 0x0001_000C",
        "LSU AXI AWADDR   : 0x0001_0008",
        "LSU WSTRB        : 8'b1111_0000",
        "MIG beat address : 0x0001_0000",
        "checker slice    : RDATA[127:96] = 32'h0000_01BC",
    ])

    add_page_break(doc)
    add_heading(doc, "5. ATG 初始化与永久总线所有权切换", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(boot_png), width=Inches(6.45))
    add_caption(doc, "图 5-1  从上电到 EH2 运行的受控启动序列")

    add_heading(doc, "5.1 DDR4-1 与 DDR4-2 的一次性 ATG", 2)
    add_table(doc, ["ATG", "启动前置条件", "写入内容", "完成后的行为"], [
        ("ATG0", "DDR4-1 calib 完成", "DDR 兼容程序映像，起始地址 0", "done 锁存；ATG0 永久保持复位；IFU 获得主通路"),
        ("ATG1", "DDR4-2 calib 完成", "0x10000~0x1000C 数据初值", "done 锁存；ATG1 永久保持复位；LSU 获得主通路"),
    ], [1250, 2050, 2850, 3210], first_col_bold=True)
    add_paragraph(doc, "两路 ATG 各自使用独立的地址、数据、掩码和控制 COE。启动时 ATG 是对应 MIG 的唯一活动主机，EH2 端口虽然存在，但所有 AXI VALID 均被所有权选择器隔离。")

    add_heading(doc, "5.2 所有权如何永久交给 EH2", 2)
    add_numbers(doc, [
        "ATG 完成信号在 100 MHz ATG 域置位 done_latched。该寄存器只在板级系统复位时清零。",
        "done_latched 分别经过两级同步器进入对应 MIG UI 时钟域，形成 done_sync。",
        "axi_owner_mux2 使用 select_b = done_sync：select_b=0 时选择 A 端 ATG；select_b=1 时选择 B 端 EH2。",
        "ATG 的 resetn 定义为 system_resetn && calib_sync && !done_latched。done_latched 一旦为 1，ATG resetn 持续为 0，ATG 不能再次发出 AXI 请求。",
        "由于 done_latched 不会在运行期清零，选择器不会返回 A 端。DDR4-1 的主通路永久交给 IFU，DDR4-2 的主通路永久交给 LSU，直到整板复位。",
    ])
    add_callout(doc, "不可逆性保证", "“永久”是指在一次上电周期内不可逆：ATG 完成锁存既控制 mux 选择，又强制 ATG 处于复位。即使 ATG 内部状态异常，接口也无法重新竞争总线。只有 system_resetn 重新复位整个系统后，所有权才会回到初始化阶段。", LIGHT_GREEN, GREEN)

    add_heading(doc, "5.3 结果读取器与“永久所有权”的关系", 2)
    add_paragraph(doc, "DDR4-2 的业务主通路仍然永久由 LSU 所有；结果读取器不是 ATG，也不会使 mux 回到初始化端。它作为终态验证覆盖层，在观察到标记的最终写响应后才短暂接管 DDR4-2 读通道，读取一次 512 位数据并锁存 pass/fail。该覆盖发生在程序结果已经写入以后，只服务于 LED 验证。")
    add_callout(doc, "运行期主机边界", "ATG1 只负责 DDR4-2 的一次性初值写入，完成后永久退出。处理器运行期间 DDR4-2 的唯一业务访问者是 EH2 LSU；终态读取器只在完成判定后执行一次只读检查。", LIGHT_BLUE, BLUE)

    add_page_break(doc)
    add_heading(doc, "6. EH2 硬件初始化与 ECC", 1)
    add_heading(doc, "6.1 官方 testbench 初始化方式", 2)
    add_paragraph(doc, "Cores-VeeR-EH2-main 的 testbench/tb_top.sv 使用 init_dccm 和 init_iccm 任务，通过层次化路径直接把存储宏的 39 位数据/ECC 单元赋零。这种写法可以快速建立仿真初态，但层次化 memory 赋值不是可综合硬件，不能直接用于 FPGA 上电初始化。")

    add_heading(doc, "6.2 本工程使用的硬件接口", 2)
    add_callout(doc, "使用的处理器接口：EH2 DMA AXI 从端口", "eh2_hw_init 作为外部 AXI 主机，连接 EH2 包装层暴露的 dma_axi_aw*、dma_axi_w*、dma_axi_b* 信号。它不通过 IFU 或 LSU，也不访问外部 DDR，而是利用 EH2 已有 DMA 地址译码和本地存储写通路访问 DCCM/ICCM。", LIGHT_GREEN, GREEN)
    add_table(doc, ["区域", "地址范围", "写次数", "每次事务"], [
        ("DCCM 64 KiB", "0xF0040000 ~ 0xF004FFF8", "8192", "64 bit 写零，AWSIZE=3，AWLEN=0，WSTRB=0xFF"),
        ("ICCM 64 KiB", "0xEE000000 ~ 0xEE00FFF8", "8192", "64 bit 写零，AWSIZE=3，AWLEN=0，WSTRB=0xFF"),
        ("合计", "128 KiB", "16384", "每笔检查 BID=0 且 BRESP=OKAY"),
    ], [1800, 2900, 1350, 3310], first_col_bold=True)

    add_heading(doc, "6.3 为什么能得到正确 ECC", 2)
    add_paragraph(doc, "DMA AXI 写事务进入 EH2 正常的内部存储访问路径。DCCM/ICCM 的存储宏不是简单保存 64 位数据，而是按处理器内部组织把数据写入并生成对应的 SECDED ECC 位。因此写入 WDATA=0 并不等于生硬地把未知 ECC 保留下来；EH2 的存储写逻辑会为零数据计算正确校验码。")
    add_bullets(doc, [
        "每笔为完整 64 位写，WSTRB=0xFF，避免局部写需要读取未知旧数据。",
        "写地址按 8 字节递增，覆盖两个 64 KiB 地址区间的每一个双字。",
        "状态机等待 B 响应后才发出下一笔，保证顺序和错误可检测性。",
        "任何 BID 或 BRESP 异常都会锁存 init_error，禁止处理器进入正常运行。",
    ])

    add_heading(doc, "6.4 复位、debug halt 与启动的精确定义", 2)
    add_paragraph(doc, "EH2 的 DMA AXI 接口不能在处理器 rst_l 物理拉低时工作。因此工程采用“先释放电气复位、但保持架构执行停止”的方法：两路 MIG 校准和 ATG 完成后，先释放 dbg_rst_l，再释放 rst_l；同时配置 mpc_reset_run_req=0，使核心进入 debug halt，不取指、不执行用户程序。")
    add_numbers(doc, [
        "复位计数达到 2 时释放 dbg_rst_l，达到 5 时释放核心 rst_l，保持官方 testbench 中 debug/core 复位分离的意图。",
        "等待 o_debug_mode_status 确认 EH2 已进入 debug halt。此时核心逻辑可响应 DMA，但架构执行仍被保持。",
        "eh2_hw_init 完成 16384 次写零并确认无 AXI 错误。",
        "控制逻辑脉冲 mpc_debug_run_req，并等待 mpc_debug_run_ack。",
        "rst_vec 固定为 0，EH2 从地址 0 开始取指。",
    ])
    add_callout(doc, "关于“复位解除前清零”的说明", "从用户程序可见行为看，DCCM/ICCM 在处理器开始执行前已经整体清零并具有正确 ECC；但从引脚时序看，rst_l 已先释放，EH2 处于 debug halt。这样既满足 DMA 接口工作条件，又确保初始化完成前不会执行任何指令。", LIGHT_AMBER, AMBER)

    add_page_break(doc)
    add_heading(doc, "7. 程序映像与 AMO 兼容性处理", 1)
    add_heading(doc, "7.1 外部 DDR AMO 的限制", 2)
    add_paragraph(doc, "原始 program.hex 在 PC=0x38 和 0x3C 含两条 AMO 原子指令。VeeR-EH2 RTL 的 LSU 地址检查明确把 AMO 限制在 DCCM：原子访问目标若位于外部 DDR，会产生 Store/AMO access fault。前仿已准确复现该异常；这不是 MIG 或 DDR4 的故障，而是 EH2 对外部 AXI 存储的体系结构限制。")

    add_heading(doc, "7.2 替换为普通 DDR 读写", 2)
    add_paragraph(doc, "按照用户授权，工程生成 init/program_ddr_compatible.hex；桌面的原始 program.hex 未被修改。两条 AMO 被展开为普通 load/add/store 序列，后续代码向后移动 8 字节，并调整最终 PC 相对立即数以保持目标地址 0x1000C 不变。")
    add_table(doc, ["PC", "指令", "机器码", "作用"], [
        ("0x38", "lw a6, 0(a5)", "0x0007A803", "从 DDR4-2 读取旧值"),
        ("0x3C", "add a7, a6, a1", "0x00B808B3", "在寄存器内执行普通加法"),
        ("0x40", "sw a7, 0(a5)", "0x0117A023", "把加法结果写回 DDR4-2"),
        ("0x44", "sw a3, 0(a5)", "0x00D7A023", "执行原程序后续普通写"),
        ("0x50", "auipc ...", "—", "调整后的尾部起点"),
        ("0x54", "addi ...", "0xFBC98993", "立即数 -68，保持最终目标 0x1000C"),
        ("0x58", "sw ...", "0x00A9A023", "写入最终结果 444（0x1BC）"),
    ], [1250, 2200, 1900, 4010], first_col_bold=True)

    add_heading(doc, "7.3 预期数据演算与终态", 2)
    add_table(doc, ["地址", "ATG 初值", "程序行为", "终值"], [
        ("0x10000", "25", "作为除法输入", "25"),
        ("0x10004", "4", "作为除法输入", "4"),
        ("0x10008", "100", "普通读改写，后续写入除法结果", "6"),
        ("0x1000C", "0", "压缩指令尾部计算 444 并存储", "444 / 0x000001BC"),
    ], [1700, 1800, 3900, 1960], first_col_bold=True)
    add_paragraph(doc, "最终 store 在 PC=0x58 退休。地址 0x5C 的程序映像为 0，EH2 随后产生 illegal instruction（cause=2）并按当前 trap 配置回到地址 0；该异常发生在正确结果已经写入以后，不影响终态检查。若需要处理器在完成后静止，可在后续软件版本中加入显式 WFI 或自循环。")

    add_heading(doc, "7.4 LED 判定", 2)
    add_table(doc, ["LED", "置位条件", "仿真结果"], [
        ("LED0", "两个 ATG 均 done 且无 ATG error", "1"),
        ("LED1", "观察到至少一次 EH2 IFU AXI 取指握手", "1"),
        ("LED2", "观察到至少一次 EH2 LSU AXI 写事务", "1"),
        ("LED3", "TCM scrub 无错误，且 DDR4-2 终态读取为 0x000001BC", "1"),
    ], [1300, 5550, 2510], first_col_bold=True)

    add_page_break(doc)
    add_heading(doc, "8. 验证结果", 1)
    add_heading(doc, "8.1 系统级前仿", 2)
    add_paragraph(doc, "系统 testbench 使用两组 512 位 AXI DDR 模型和生产参数的完整 TCM 扫描范围，验证双 ATG 装载、所有权永久交接、64 KiB DCCM 与 64 KiB ICCM 清零、EH2 50 MHz 运行、IFU 取指、LSU 普通 DDR 读写、终态读取和四颗 LED。")
    add_code_box(doc, [
        "PASS: dual ATG load, permanent EH2 handover, 50 MHz execution,",
        "      DDR4-2 result 0x000001BC, LED=1111",
        "PASS: production EH2 scrub issued 16384 ordered 64-bit zero writes",
        "      across full DCCM and ICCM ranges",
    ])
    add_table(doc, ["检查点", "结果", "证据"], [
        ("双 ATG 初始化", "通过", "程序与数据初值写入完成"),
        ("ATG 永久退出", "通过", "done 后复位保持，mux 不返回 A 端"),
        ("50 MHz EH2 退休轨迹", "通过", "PC 0x00 至 0x58 正常退休"),
        ("IFU AXI 取指", "通过", "LED1 置位，首个 AR 地址为 0"),
        ("LSU AXI 写", "通过", "LED2 置位，B 响应成功"),
        ("最终 DDR 数据", "通过", "0x1000C = 0x000001BC"),
        ("完整 DCCM/ICCM 清零", "通过", "16384 次有序 64 位 DMA 写"),
    ], [2550, 1450, 5360], first_col_bold=True)

    add_heading(doc, "8.2 综合后功能仿真", 2)
    add_paragraph(doc, "综合后 testbench 只驱动顶层差分时钟和复位并观察顶层 LED，不使用层次 force。ATG、永久所有权 mux、Xilinx AXI Clock Converter、Data Width Converter 和 EH2 均来自综合网表；物理 MIG/DDR PHY 由可综合 AXI DDR 模型替代，因此该用例验证处理器到 MIG AXI 边界的网表级通路，不替代板级 DDR PHY 校准与信号完整性测试。为控制 XSim 网表仿真时间，本用例把两块 TCM 的扫描各缩短为 32 次写；完整 128 KiB 覆盖已由系统级生产参数前仿验证。")
    add_table(doc, ["综合后检查点", "结果"], [
        ("双 MIG 模型与双 ATG 完成", "通过"),
        ("DMA TCM scrub 与 debug resume", "通过"),
        ("EH2 从地址 0 发出 IFU 取指", "通过"),
        ("EH2 发出 LSU 写事务", "通过"),
        ("DDR 终态回读与 LED=1111", "通过"),
    ], [5200, 4160], first_col_bold=True)

    add_heading(doc, "8.3 综合与实现时序", 2)
    add_table(doc, ["阶段", "WNS", "TNS", "WHS", "THS", "结论"], [
        ("综合后", "+0.893 ns", "0", "—", "—", "满足"),
        ("布局布线后", "+0.202 ns", "0", "+0.010 ns", "0", "满足"),
        ("EH2 core_clk 域", "+6.486 ns", "0", "+0.012 ns", "0", "满足"),
        ("脉宽", "+0.046 ns", "0", "—", "—", "满足"),
    ], [1850, 1500, 1300, 1500, 1300, 1910], first_col_bold=True)
    add_callout(doc, "时序结论", "Vivado 实现报告明确给出“All user specified timing constraints are met”。布局布线后 setup、hold 和 pulse-width 均无失败端点。", LIGHT_GREEN, GREEN)

    add_heading(doc, "8.4 DDR 总线偏斜与路由", 2)
    add_table(doc, ["项目", "数量 / 最差值", "状态"], [
        ("post-route bus skew checks", "44 项；最小裕量 2.367 ns", "全部 MET"),
        ("routable nets", "211,108", "全部完成"),
        ("fully routed nets", "211,108", "完成"),
        ("nets with routing errors", "0", "无错误"),
    ], [3300, 3500, 2560], first_col_bold=True)

    add_heading(doc, "8.5 时钟频率汇总", 2)
    add_table(doc, ["时钟", "实现频率", "用途"], [
        ("core_clk", "50.000 MHz", "VeeR-EH2、硬件初始化状态机；GCCLKT0 BY44/CA44"),
        ("atg_clk", "100.000 MHz", "双 ATG 命令生成"),
        ("c0_sys_clk", "76.150 MHz", "DDR4-1 MIG 参考输入"),
        ("c1_sys_clk", "76.150 MHz", "DDR4-2 MIG 参考输入"),
        ("MIG0 mmcm_clkout0", "266.525 MHz", "DDR4-1 AXI UI"),
        ("MIG1 mmcm_clkout0_1", "266.525 MHz", "DDR4-2 AXI UI"),
        ("debug TCK（内部）", "20.000 MHz", "Vivado debug hub"),
    ], [2900, 2300, 4160], first_col_bold=True)

    add_page_break(doc)
    add_heading(doc, "9. 实现质量、资源与功耗", 1)
    add_heading(doc, "9.1 资源利用率", 2)
    add_table(doc, ["资源", "已用", "可用", "利用率"], [
        ("CLB LUT", "115,282", "4,085,760", "2.82%"),
        ("LUT as logic", "106,028", "4,085,760", "2.59%"),
        ("LUT memory", "9,254", "—", "0.97%（对应可用 LUTRAM）"),
        ("Registers", "98,325", "8,171,520", "1.20%"),
        ("Block RAM tile", "71", "2,160", "3.29%"),
        ("URAM", "4", "320", "1.25%"),
        ("DSP", "10", "3,840", "0.26%"),
        ("Bonded IOB", "267", "2,072", "12.89%"),
        ("BUFGCE / PLL / MMCM", "12 / 6 / 2", "—", "—"),
    ], [2700, 1900, 2200, 2560], first_col_bold=True)
    add_paragraph(doc, "主要资源来自两组 MIG、四个 AXI 转换链路和 VeeR-EH2 本体。器件容量裕量充足，BRAM/URAM 和 DSP 利用率均较低。")

    add_heading(doc, "9.2 CDC 报告", 2)
    add_table(doc, ["类别", "数量", "解释"], [
        ("CDC-3 Info", "408", "单比特 ASYNC_REG 同步器，符合预期"),
        ("CDC-9 Info", "64", "异步复位同步释放结构，符合预期"),
        ("CDC-15 Warning", "6", "均位于 Xilinx IP：2 个 AXI 异步 FIFO、4 个 debug_hub FIFO"),
        ("Critical CDC", "0", "无用户顶层关键 CDC"),
    ], [2450, 1400, 5510], first_col_bold=True)
    add_paragraph(doc, "此前把多个异步状态先组合再同步会造成 CDC 风险，最终设计改为六个状态独立两级同步、在目标域组合。实现报告中的 6 个 warning 均是工具已知的 Xilinx IP 内部异步 FIFO 结构。")

    add_heading(doc, "9.3 DRC 与方法学警告", 2)
    add_table(doc, ["报告", "错误 / 严重警告", "普通警告 / 建议", "处置结论"], [
        ("Implementation DRC", "0 / 0", "52 warnings + 2 advisories", "不阻塞；设计 Fully Routed"),
        ("Methodology", "0 / 0", "217 warnings", "均为优化或约束建议"),
    ], [2300, 2150, 2580, 2330], first_col_bold=True)
    add_bullets(doc, [
        "DPIP/DPOP/DPOR：EH2 内部 DSP 管线或异步复位的优化建议；最终时序已满足。",
        "PDCN-1569：Vivado debug_hub LUT 方程项，不属于用户业务逻辑。",
        "RTSTAT-10：MIG 校准/debug 内部无可路由负载网络，不影响 DDR 数据通路。",
        "TIMING-18/24：外部开关和 LED 不是同步采样数据接口，未设置输入/输出延迟；MIG 使用 IP 生成约束覆盖接口时序。",
        "SECHK-3/4：高速收发器和特殊 I/O 数量的器件建议，当前用量低于阈值。",
    ])

    add_heading(doc, "9.4 功耗估算", 2)
    add_table(doc, ["指标", "Vivado 估算"], [
        ("总片上功耗", "8.874 W"),
        ("动态功耗", "2.968 W"),
        ("静态功耗", "5.906 W"),
        ("估算结温", "28.9 °C"),
        ("置信度", "Low"),
    ], [3500, 5860], first_col_bold=True)
    add_callout(doc, "功耗使用限制", "该结果为未导入实测活动率的向量无关估算，且部分 XPHY/VREF 模型不可用，不能替代上板电流、结温与散热测试。ATG 完成后长期保持复位会使高扇出复位活动率假设偏保守。", LIGHT_AMBER, AMBER)

    add_page_break(doc)
    add_heading(doc, "10. 交付物与复现", 1)
    add_heading(doc, "10.1 主要工程文件", 2)
    add_table(doc, ["类别", "位置", "说明"], [
        ("Vivado 工程", "build/vivado/eh2_dual_ddr.xpr", "可重新打开工程和实现运行"),
        ("最终比特流", "output/bitstreams/eh2_dual_ddr_gclkt0_50mhz_led4.bit", "GCCLKT0=50 MHz 版本"),
        ("实现后 DCP", "reports/eh2_dual_ddr_impl.dcp", "实现状态检查与后续分析"),
        ("顶层 RTL", "rtl/eh2_dual_ddr_top.sv", "系统连接、复位和 LED"),
        ("EH2 包装", "rtl/eh2_core_wrapper_hw.sv", "处理器、debug 与 AXI 接口"),
        ("TCM 初始化", "rtl/eh2_hw_init.sv", "DMA AXI 16384 次写零"),
        ("总线所有权", "rtl/axi_owner_mux2.sv", "ATG/EH2 二选一主机 mux"),
        ("最终结果检查", "rtl/ddr_result_checker.sv", "DDR4-2 终态读回"),
        ("板级约束", "constraints/eh2_dual_ddr_v19p.xdc", "时钟、DDR、复位和 LED"),
        ("程序映像", "init/program_ddr_compatible.hex", "AMO 展开后的 DDR 版本"),
        ("系统仿真日志", "reports/system_pre_sim.log", "终值和 LED PASS"),
        ("TCM 仿真日志", "reports/hw_init_pre_sim.log", "完整 128 KiB scrub PASS"),
        ("实现报告", "reports/*_impl.rpt", "timing、bus skew、CDC、DRC、power、utilization"),
    ], [1900, 4530, 2930], first_col_bold=True)

    add_heading(doc, "10.2 复现流程", 2)
    add_numbers(doc, [
        "使用 Vivado 2023.2 打开 build/vivado/eh2_dual_ddr.xpr，确认器件为 xcvu19p_CIV-fsva3824-1-e。",
        "运行 scripts/run_pre_sim.tcl 和 scripts/run_hw_init_sim.tcl，确认两个 PASS 日志。",
        "运行 scripts/run_synthesis.tcl，检查综合后 timing/CDC/utilization 报告。",
        "运行 scripts/run_implementation.tcl，生成实现报告、DCP 和 BIT。",
        "下载比特流前确认板卡两块 SODIMM 型号、供电和复位开关状态符合手册。",
    ])

    add_heading(doc, "10.3 上板验证建议", 2)
    add_numbers(doc, [
        "上电后先观察两路 MIG 校准状态（必要时通过 ILA 标记），确认两路 init_calib_complete 都置位。",
        "在板卡时钟配置中使能 GCCLKT0 并设为 50 MHz，确认 TOP_GCLK0 选择 SI5338 输出而不是 SCLK2 来源。",
        "确认 LED0 点亮，表示双 ATG 完成且无错误；随后 ATG 应持续处于复位。",
        "确认 LED1 在处理器产生第一次 IFU AXI 取指握手后锁存点亮。",
        "确认 LED2 在处理器发出 LSU AXI 写操作后锁存点亮。",
        "确认 LED3 最终点亮，表示 TCM 清零无错误且 DDR4-2 地址 0x1000C 读回 0x000001BC。",
        "若 LED0 不亮，优先检查 GCLK3、DDR 管脚、SODIMM 型号和 MIG 校准；若 LED1 不亮，检查 GCCLKT0、EH2 复位/debug run 和 IFU 通路；若 LED2 不亮，检查 LSU/BRESP；若 LED3 不亮，检查 RRESP 和 checker 状态。",
        "上板采集电源轨电流和器件温度，用实测活动率重新评估功耗与散热裕量。",
    ])

    add_heading(doc, "10.4 参考来源", 2)
    add_bullets(doc, [
        r"VeriTiger-V19P-A14 用户手册 CN V13（2024-11-29）：E:\VeriTiger-V19P-A14\03-使用手册\VeriTiger-V19P-A14用户手册_CN_V13_20241129-1.pdf",
        r"MAC 参考工程：D:\eh2_fpga\mac_fifo_dma_proj",
        r"VeeR-EH2 RTL 与 testbench：D:\eh2_fpga\Cores-VeeR-EH2-main\Cores-VeeR-EH2-main",
        r"原始程序映像：C:\Users\18217\Desktop\program.hex（保持未修改）",
        r"本工程：D:\eh2_fpga\eh2_veri_iss_proj",
    ])

    add_callout(doc, "最终结论", "在仿真和静态实现范围内，双 DDR4、一次性 ATG、不可逆总线交接、EH2 DMA TCM/ECC 初始化、50 MHz 程序执行和 0x1BC 终态判定均已验证；比特流成功生成且实现时序满足。剩余工作仅为实体板卡下载与实验室级硬件确认。", LIGHT_GREEN, GREEN)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
